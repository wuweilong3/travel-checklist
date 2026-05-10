from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

def generate_report():
    doc = Document()
    
    # 设置页面
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # 封面
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("项目中期报告")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("旅睿（TravelWise）- 智能旅行清单助手")
    subtitle_run.font.size = Pt(16)
    
    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.cell(0, 0).text = "课程：智能应用项目开发综合实践"
    info_table.cell(1, 0).text = "作业名称：项目中期报告"
    info_table.cell(2, 0).text = "总分：100分"
    info_table.cell(3, 0).text = "提交日期：2026年5月"
    
    doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 第一章
    doc.add_heading("一、作业概述", level=1)
    doc.add_heading("1.1 作业目的", level=2)
    doc.add_paragraph("1. 检查项目按计划执行的实际进展")
    doc.add_paragraph("2. 验证核心功能是否已实现并可运行")
    doc.add_paragraph("3. 发现问题并及时调整，避免期末突击")
    doc.add_paragraph("4. 督促持续推进项目，保持开发节奏")
    doc.add_paragraph("5. 总结与编程Agent协作的经验，提升协同开发能力")
    
    doc.add_heading("1.2 与第一次汇报的关系", level=2)
    compare_table = doc.add_table(rows=6, cols=3)
    compare_table.style = 'Table Grid'
    compare_table.cell(0, 0).text = "对比项"
    compare_table.cell(0, 1).text = "第一次汇报"
    compare_table.cell(0, 2).text = "中期报告"
    compare_table.cell(1, 0).text = "阶段"
    compare_table.cell(1, 1).text = "规划阶段"
    compare_table.cell(1, 2).text = "执行阶段"
    compare_table.cell(2, 0).text = "重点"
    compare_table.cell(2, 1).text = "做什么、为什么做"
    compare_table.cell(2, 2).text = "做了什么、做得怎样"
    compare_table.cell(3, 0).text = "产出"
    compare_table.cell(3, 1).text = "需求文档、设计文档"
    compare_table.cell(3, 2).text = "代码、运行效果、测试结果"
    compare_table.cell(4, 0).text = "验证"
    compare_table.cell(4, 1).text = "需求自洽、设计可行"
    compare_table.cell(4, 2).text = "必须可运行、可演示"
    compare_table.cell(5, 0).text = "风险"
    compare_table.cell(5, 1).text = "方向偏差"
    compare_table.cell(5, 2).text = "进度滞后、技术困难"
    
    doc.add_paragraph("核心区别：中期报告是'做出来的'，必须有真实运行效果")
    
    doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 第二章
    doc.add_heading("二、项目进展总览", level=1)
    doc.add_heading("2.1 项目基本信息", level=2)
    doc.add_paragraph("• 项目名称：旅睿（TravelWise）- 智能旅行清单助手")
    doc.add_paragraph("• 目标用户：计划旅行的人群")
    doc.add_paragraph("• 核心问题：解决旅行清单准备问题")
    doc.add_paragraph("• 核心价值：根据目的地、季节、人群自动生成场景化清单")
    
    doc.add_heading("2.2 迭代计划执行情况", level=2)
    plan_table = doc.add_table(rows=4, cols=4)
    plan_table.style = 'Table Grid'
    plan_table.cell(0, 0).text = "周次"
    plan_table.cell(0, 1).text = "计划内容"
    plan_table.cell(0, 2).text = "完成状态"
    plan_table.cell(0, 3).text = "说明"
    plan_table.cell(1, 0).text = "10周"
    plan_table.cell(1, 1).text = "需求分析与架构设计"
    plan_table.cell(1, 2).text = "✅ 已完成"
    plan_table.cell(1, 3).text = "完成需求文档和技术方案"
    plan_table.cell(2, 0).text = "11周"
    plan_table.cell(2, 1).text = "后端API开发与LLM集成"
    plan_table.cell(2, 2).text = "✅ 已完成"
    plan_table.cell(2, 3).text = "实现核心Agent逻辑"
    plan_table.cell(3, 0).text = "12周"
    plan_table.cell(3, 1).text = "前端界面开发"
    plan_table.cell(3, 2).text = "✅ 已完成"
    plan_table.cell(3, 3).text = "React聊天界面和清单视图"
    
    doc.add_heading("2.3 总体完成度", level=2)
    doc.add_paragraph("总体完成度：90%")
    doc.add_paragraph("已完成功能：智能聊天界面、旅行清单生成、清单完整性检查、物品增删改查、饮食偏好推荐、用户习惯记忆")
    doc.add_paragraph("剩余工作：主动服务推送、移动端适配")
    
    doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 第三章
    doc.add_heading("三、已完成功能详述", level=1)
    doc.add_heading("3.1 已完成功能列表", level=2)
    func_table = doc.add_table(rows=8, cols=4)
    func_table.style = 'Table Grid'
    func_table.cell(0, 0).text = "功能名称"
    func_table.cell(0, 1).text = "实现周次"
    func_table.cell(0, 2).text = "核心代码位置"
    func_table.cell(0, 3).text = "运行效果"
    func_table.cell(1, 0).text = "智能聊天界面"
    func_table.cell(1, 1).text = "12周"
    func_table.cell(1, 2).text = "frontend/src/App.tsx"
    func_table.cell(1, 3).text = "✅ 正常运行"
    func_table.cell(2, 0).text = "意图识别"
    func_table.cell(2, 1).text = "11周"
    func_table.cell(2, 2).text = "backend/agents/core.py"
    func_table.cell(2, 3).text = "✅ 正常运行"
    func_table.cell(3, 0).text = "旅行清单生成"
    func_table.cell(3, 1).text = "11周"
    func_table.cell(3, 2).text = "backend/agents/tools.py"
    func_table.cell(3, 3).text = "✅ 正常运行"
    func_table.cell(4, 0).text = "清单完整性检查"
    func_table.cell(4, 1).text = "11周"
    func_table.cell(4, 2).text = "backend/agents/tools.py"
    func_table.cell(4, 3).text = "✅ 正常运行"
    func_table.cell(5, 0).text = "物品增删改查"
    func_table.cell(5, 1).text = "12周"
    func_table.cell(5, 2).text = "backend/api/routes.py"
    func_table.cell(5, 3).text = "✅ 正常运行"
    func_table.cell(6, 0).text = "饮食偏好推荐"
    func_table.cell(6, 1).text = "12周"
    func_table.cell(6, 2).text = "backend/agents/core.py"
    func_table.cell(6, 3).text = "✅ 正常运行"
    func_table.cell(7, 0).text = "用户习惯记忆"
    func_table.cell(7, 1).text = "11周"
    func_table.cell(7, 2).text = "backend/agents/tools.py"
    func_table.cell(7, 3).text = "✅ 正常运行"
    
    doc.add_heading("3.2 核心功能演示", level=2)
    doc.add_heading("功能1：智能清单生成", level=3)
    doc.add_paragraph("用户输入：我要去成都旅游，3天")
    doc.add_paragraph("系统响应：已为你生成【成都3日游】的清单！共15件物品，分为5大类")
    
    doc.add_heading("功能2：完整性检查", level=3)
    doc.add_paragraph("用户输入：检查一下有没有遗漏")
    doc.add_paragraph("系统响应：清单完整度检查完成！当前完整度：70%")
    
    doc.add_heading("功能3：饮食偏好推荐", level=3)
    doc.add_paragraph("用户输入：我喜欢吃辣，请给我推荐合适的旅游地点")
    doc.add_paragraph("系统响应：根据你喜欢吃辣的口味，推荐：成都、重庆、长沙")
    
    doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 第四章
    doc.add_heading("四、技术实现与架构", level=1)
    doc.add_heading("4.1 实际技术架构", level=2)
    doc.add_paragraph("前端(React) → 后端(FastAPI) → DashScope LLM → JSON存储")
    
    doc.add_heading("4.2 关键技术决策", level=2)
    tech_table = doc.add_table(rows=4, cols=3)
    tech_table.style = 'Table Grid'
    tech_table.cell(0, 0).text = "决策项"
    tech_table.cell(0, 1).text = "选择方案"
    tech_table.cell(0, 2).text = "理由"
    tech_table.cell(1, 0).text = "LLM引擎"
    tech_table.cell(1, 1).text = "阿里云百炼DashScope"
    tech_table.cell(1, 2).text = "国内访问稳定"
    tech_table.cell(2, 0).text = "后端框架"
    tech_table.cell(2, 1).text = "FastAPI"
    tech_table.cell(2, 2).text = "高性能异步框架"
    tech_table.cell(3, 0).text = "前端框架"
    tech_table.cell(3, 1).text = "React + TypeScript"
    tech_table.cell(3, 2).text = "组件化开发，类型安全"
    
    doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 第五章
    doc.add_heading("五、与Agent协作总结", level=1)
    doc.add_heading("5.1 协同方式", level=2)
    doc.add_paragraph("工具：Trae IDE + Claude 3.5")
    doc.add_paragraph("模式：需求分析 → 方案讨论 → 代码实现 → 测试验证")
    
    doc.add_heading("5.2 协同案例", level=2)
    doc.add_paragraph("案例1：意图识别问题修复 - 识别准确率从60%提升至95%")
    doc.add_paragraph("案例2：饮食偏好推荐功能 - 支持8种饮食偏好类型")
    
    doc.add_heading("5.3 协同经验", level=2)
    doc.add_paragraph("1. 精确指令能大幅提高Agent输出质量")
    doc.add_paragraph("2. 分步验证的重要性：先验证API，再集成前端")
    doc.add_paragraph("3. 接受不完美的初始方案，逐步优化")
    
    doc.add_section(WD_SECTION.NEW_PAGE)
    
    # 第六章
    doc.add_heading("六、问题与调整", level=1)
    doc.add_heading("6.1 遇到的问题", level=2)
    doc.add_paragraph("• LLM API调用失败 → 实现mock模式自动降级")
    doc.add_paragraph("• 中文编码乱码 → 设置UTF-8编码")
    doc.add_paragraph("• 意图识别不完善 → 迭代优化正则匹配规则")
    
    doc.add_heading("6.2 计划调整", level=2)
    doc.add_paragraph("• 第13周：优先完成主动服务功能")
    doc.add_paragraph("• 第14周：用户体验优化，包括移动端适配")
    
    # 第七章
    doc.add_heading("七、后续计划", level=1)
    doc.add_heading("7.1 剩余功能规划", level=2)
    doc.add_paragraph("• 主动服务推送（高优先级，第13周）")
    doc.add_paragraph("• 移动端适配（高优先级，第13周）")
    doc.add_paragraph("• 清单导出功能（中优先级，第14周）")
    
    doc.add_heading("7.2 期末目标", level=2)
    doc.add_paragraph("• 功能完整度：100%")
    doc.add_paragraph("• 部署至公网，提供在线演示地址")
    
    # 保存文档
    doc.save("项目中期报告.docx")
    print("Word文档已生成：项目中期报告.docx")

if __name__ == "__main__":
    generate_report()